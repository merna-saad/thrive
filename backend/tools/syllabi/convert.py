"""Convert the raw syllabus corpus to markdown and index its filenames.

    python backend/tools/syllabi/convert.py

The PDFs under `backend/data/syllabi/` are the source of truth. They get
reconverted whenever a new quarter's syllabi land. Everything this script
writes under `backend/data/syllabi_md/` is a derived artifact: markdown
mirroring the source folders, plus `index.json` and `report.md`.

This pass is conversion and filename parsing only. No LLM calls, no skills
extraction, no database models.

Locations are configuration -- see --source and --out. Paths recorded in
index.json are repo-relative, never absolute.
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field, fields
from pathlib import Path

import pymupdf
import pymupdf4llm
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).resolve().parent))

from course_index import build_courses, build_title_review
from filename_parser import Offering, ParseError, parse_filename

# The offering fields are flattened into each index record, so rebuilding a
# Record from JSON means knowing which keys came from the parser.
OFFERING_FIELDS = tuple(f.name for f in fields(Offering))

REPO_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_SOURCE = Path("backend/data/syllabi")
DEFAULT_OUT = Path("backend/data/syllabi_md")

SUPPORTED_SUFFIXES = {".pdf", ".docx"}

# Folders holding documents that are not course offerings. Files here are
# converted but never run through the offering parser, so they do not show up
# as parse failures.
CATALOG_DIRS = {"Course Descriptions"}

# Below this many extracted characters a syllabus is almost certainly a scan
# or an image-only PDF, and the markdown is not usable downstream.
THIN_TEXT_CHARS = 500


@dataclass(frozen=True)
class Roots:
    """Where things live. Absolute for I/O, repo-relative for the record."""

    source: Path
    out: Path
    rel_source: Path
    rel_out: Path


@dataclass
class Record:
    """One source file: what it is, where it went, what parsed out of it."""

    kind: str  # "offering" | "catalog" | "unparsed"
    filename: str
    source_path: str
    markdown_path: str | None
    converted: bool
    characters: int
    pages: int = 0
    pages_without_text_layer: int = 0
    conversion_error: str | None = None
    parse_error: str | None = None
    offering: dict[str, str | None] | None = None

    @classmethod
    def from_dict(cls, data: dict) -> Record:
        """Inverse of `to_dict`, so --report-only can skip reconversion.

        A record only carries an offering if EVERY offering field is present.
        Testing for any single key is not enough: `filename` sits on both the
        record and the offering, so a catalog record would otherwise rebuild
        as an offering holding nothing but its filename.
        """
        offering = (
            {k: data[k] for k in OFFERING_FIELDS} if all(k in data for k in OFFERING_FIELDS) else {}
        )
        return cls(
            kind=data["kind"],
            filename=data["filename"],
            source_path=data["source_path"],
            markdown_path=data["markdown_path"],
            converted=data["converted"],
            characters=data["characters"],
            pages=data.get("pages", 0),
            pages_without_text_layer=data.get("pages_without_text_layer", 0),
            conversion_error=data.get("conversion_error"),
            parse_error=data.get("parse_error"),
            offering=offering or None,
        )

    def to_dict(self) -> dict:
        base = {
            "kind": self.kind,
            "filename": self.filename,
            "source_path": self.source_path,
            "markdown_path": self.markdown_path,
            "converted": self.converted,
            "characters": self.characters,
            "pages": self.pages,
            "pages_without_text_layer": self.pages_without_text_layer,
        }
        if self.offering:
            base.update(self.offering)
        if self.conversion_error:
            base["conversion_error"] = self.conversion_error
        if self.parse_error:
            base["parse_error"] = self.parse_error
        return base


@dataclass
class Totals:
    records: list[Record] = field(default_factory=list)

    def of_kind(self, kind: str) -> list[Record]:
        return [r for r in self.records if r.kind == kind]


def _capturing_stdout(fn, *args, **kwargs) -> tuple[object, str]:
    """Call `fn`, returning its result and everything written to fd 1.

    `contextlib.redirect_stdout` is not enough: the parser chatter comes out
    of the C layer, straight to the file descriptor, past `sys.stdout`.
    """
    sys.stdout.flush()
    saved = os.dup(1)
    with tempfile.TemporaryFile() as sink:
        try:
            os.dup2(sink.fileno(), 1)
            result = fn(*args, **kwargs)
        finally:
            os.dup2(saved, 1)
            os.close(saved)
        sink.seek(0)
        return result, sink.read().decode("utf-8", "replace")


def text_layer_stats(path: Path) -> tuple[int, int]:
    """(pages, pages with no extractable text layer).

    Measured with plain PyMuPDF, before pymupdf4llm gets involved. This is the
    honest quality signal. The OCR progress the parser prints is *not*: it
    OCRs every page whether or not one has a text layer, so counting its
    chatter would flag perfectly good digital PDFs.
    """
    with pymupdf.open(path) as doc:
        return doc.page_count, sum(1 for page in doc if not page.get_text().strip())


def pdf_to_markdown(path: Path, *, use_ocr: bool = False) -> str:
    """Extract markdown from a PDF.

    pymupdf4llm reconstructs headings, list structure, and pipe tables from
    PyMuPDF's layout analysis, and orders multi-column pages by column rather
    than by raw text-block order. That matters here: syllabi are full of
    two-column schedules and grading tables.

    OCR is off by default. pymupdf-layout runs Tesseract over *every* page
    when it is on, text layer or not, and on this corpus -- where every page
    has a text layer -- the only thing it adds is a mangled transcription of
    the Rady logo ("Ray| schoolof m") wrapped in picture-text comments, in 119
    of 141 files. Tables come out identical either way. Turn it back on with
    --ocr if a future quarter brings scans; the report's text-layer count is
    what tells you.

    pymupdf-layout also prints parser chatter straight to stdout, which would
    corrupt the report this script prints. Hence the fd capture.
    """
    text, _ = _capturing_stdout(
        pymupdf4llm.to_markdown, str(path), show_progress=False, use_ocr=use_ocr
    )
    return text


def _iter_docx_blocks(document: Document):
    """Yield paragraphs and tables in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate flat lists,
    which loses the interleaving. Walking the body XML keeps it.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_to_markdown(table: Table) -> str:
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header, *body = rows
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


def docx_to_markdown(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Table):
            table = _table_to_markdown(block)
            if table:
                chunks.append(table)
            continue
        text = block.text.strip()
        if not text:
            continue
        style = (block.style.name or "") if block.style else ""
        if style.startswith("Heading"):
            level = style.replace("Heading", "").strip()
            depth = int(level) if level.isdigit() else 1
            chunks.append(f"{'#' * min(depth, 6)} {text}")
        elif style == "Title":
            chunks.append(f"# {text}")
        elif style.startswith("List"):
            chunks.append(f"- {text}")
        else:
            chunks.append(text)
    return "\n\n".join(chunks) + "\n"


def source_files(source: Path) -> list[Path]:
    files = [
        p
        for p in sorted(source.rglob("*"))
        if p.is_file() and not p.name.startswith(".") and p.suffix.lower() in SUPPORTED_SUFFIXES
    ]
    return files


def is_catalog(path: Path, source: Path) -> bool:
    return bool(CATALOG_DIRS.intersection(path.relative_to(source).parts[:-1]))


def convert_one(path: Path, roots: Roots, *, use_ocr: bool = False) -> Record:
    """Convert one file and record what came of it.

    Paths in the record are repo-relative: this index is committed and read on
    other machines, so an absolute path in it is a bug.
    """
    relative = path.relative_to(roots.source)
    out_path = (roots.out / relative).with_suffix(".md")

    record = Record(
        kind="catalog" if is_catalog(path, roots.source) else "unparsed",
        filename=path.name,
        source_path=(roots.rel_source / relative).as_posix(),
        markdown_path=(roots.rel_out / relative).with_suffix(".md").as_posix(),
        converted=False,
        characters=0,
    )

    if record.kind != "catalog":
        try:
            record.offering = parse_filename(path.name).to_dict()
            record.kind = "offering"
        except ParseError as exc:
            record.parse_error = str(exc)

    try:
        if path.suffix.lower() == ".pdf":
            record.pages, record.pages_without_text_layer = text_layer_stats(path)
            text = pdf_to_markdown(path, use_ocr=use_ocr)
        else:
            text = docx_to_markdown(path)
    except Exception as exc:  # noqa: BLE001 - one bad file must not stop the run
        record.conversion_error = f"{type(exc).__name__}: {exc}"
        record.markdown_path = None
        return record

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(text, encoding="utf-8")
    record.converted = True
    record.characters = len(text)
    return record


def build_report(totals: Totals, source: Path, out: Path, *, use_ocr: bool) -> str:
    records = totals.records
    offerings = totals.of_kind("offering")
    catalogs = totals.of_kind("catalog")
    unparsed = totals.of_kind("unparsed")

    courses = build_courses([r.offering for r in offerings])
    review = build_title_review(courses)
    multi_code = [c for c in courses if len(c.codes) > 1]
    multi_offering = [c for c in courses if c.offering_count > 1]

    failed = [r for r in records if not r.converted]
    thin = [r for r in records if r.converted and r.characters < THIN_TEXT_CHARS]
    scanned = [r for r in records if r.pages_without_text_layer]

    codes = Counter(r.offering["course_code"] for r in offerings)
    terms = Counter(r.offering["term"] for r in offerings)
    qualifiers = Counter(r.offering["qualifier"] for r in offerings if r.offering["qualifier"])
    no_instructor = [r for r in offerings if not r.offering["instructor"]]

    by_course: dict[str, list[Record]] = defaultdict(list)
    for r in offerings:
        by_course[r.offering["course_key"]].append(r)

    lines: list[str] = []
    add = lines.append

    add("# Syllabus conversion report")
    add("")
    add(f"Source: `{source.as_posix()}`  ")
    add(f"Output: `{out.as_posix()}`  ")
    add("Extractor: pymupdf4llm (PyMuPDF) for PDF, python-docx for DOCX  ")
    if use_ocr:
        add("OCR: on (`--ocr`). Tesseract runs over every page, text layer or not.")
    else:
        add("OCR: off. Every page below has a text layer, so it would only add noise.")
    add("")

    add("## Totals")
    add("")
    add(f"- Files found: {len(records)}")
    add(f"- Converted: {sum(1 for r in records if r.converted)}")
    add(f"- Failed to convert: {len(failed)}")
    add(f"- Parsed as offerings: {len(offerings)}")
    add(f"- Catalogs (not offerings, conversion only): {len(catalogs)}")
    add(f"- Filenames that did not parse: {len(unparsed)}")
    add(f"- Distinct course codes: {len(codes)}")
    add(f"- **Distinct courses (code + title): {len(courses)}**")
    add(f"- Courses offered more than once: {len(multi_offering)}")
    add(f"- Courses that have appeared under more than one code: {len(multi_code)}")
    add(f"- Title pairs awaiting manual review: {len(review)}")
    add(f"- Distinct terms: {len(terms)}")
    add(f"- Files with at least one page lacking a text layer: {len(scanned)}")
    add("")

    add("## Failed to convert")
    add("")
    if failed:
        for r in failed:
            add(f"- `{r.source_path}` — {r.conversion_error}")
    else:
        add("None.")
    add("")

    add("## Filenames that did not parse")
    add("")
    if unparsed:
        for r in unparsed:
            add(f"- `{r.filename}` — {r.parse_error}")
    else:
        add("None.")
    add("")

    add("## Thin extractions (likely scanned or image-only)")
    add("")
    add(f"Fewer than {THIN_TEXT_CHARS} characters of text.")
    add("")
    if thin:
        for r in sorted(thin, key=lambda r: r.characters):
            add(f"- `{r.source_path}` — {r.characters} chars")
    else:
        add("None.")
    add("")

    add("## Pages with no text layer")
    add("")
    add("Scanned or image-only pages. Measured with plain PyMuPDF, before extraction.")
    add("Anything listed here needs a rerun with --ocr.")
    add("")
    if scanned:
        for r in sorted(scanned, key=lambda r: -r.pages_without_text_layer):
            add(f"- `{r.source_path}` — {r.pages_without_text_layer} of {r.pages} page(s)")
    else:
        add("None.")
    add("")

    add("## Terms")
    add("")
    for term, count in sorted(terms.items()):
        add(f"- {term}: {count}")
    add("")

    add("## Trailing qualifiers")
    add("")
    if qualifiers:
        for qualifier, count in sorted(qualifiers.items()):
            add(f"- {qualifier}: {count}")
    else:
        add("None.")
    add("")

    add("## Offerings with no instructor in the filename")
    add("")
    if no_instructor:
        for r in no_instructor:
            add(f"- `{r.filename}`")
    else:
        add("None.")
    add("")

    add("## Courses with more than one offering")
    add("")
    add(f"{len(multi_offering)} of {len(courses)} courses recur. Terms are chronological.")
    add("")
    for course in multi_offering:
        add(f"- **{' / '.join(course.codes)}** — {course.title} ({course.offering_count})")
        for r in sorted(by_course[course.key], key=lambda r: r.offering["term_sort"]):
            o = r.offering
            who = o["instructor"] or "—"
            extra = f" [{o['qualifier']}]" if o["qualifier"] else ""
            add(f"  - {o['term']} · {o['course_code']} · {who}{extra} · {o['title']}")
    add("")

    add("## Courses that have appeared under more than one code")
    add("")
    add("The key contains the base code, so this can only ever be section variants")
    add("(`403` and `403R`). A course that was *renumbered* keeps its title but changes")
    add("base code, which makes it two keys — see `same_title_different_code` below.")
    add("")
    if multi_code:
        for course in multi_code:
            add(f"- **{course.title}** — {', '.join(course.codes)} ({', '.join(course.terms)})")
    else:
        add("None.")
    add("")

    add("## Title review list")
    add("")
    add("Course pairs that may be the same course. **Nothing here has been merged.**")
    add("Grouping is on exact normalised titles; these are for manual confirmation.")
    add("")
    if review:
        for reason in ("same_title_different_code", "near_identical_title", "same_code_different_title"):
            group = [p for p in review if p.reason == reason]
            if not group:
                continue
            add(f"### {reason} ({len(group)})")
            add("")
            for pair in group:
                add(f"- similarity {pair.similarity:.3f}")
                add(f"  - `{'/'.join(pair.left_codes)}` {pair.left_title!r} ({', '.join(pair.left_terms)})")
                add(f"  - `{'/'.join(pair.right_codes)}` {pair.right_title!r} ({', '.join(pair.right_terms)})")
            add("")
    else:
        add("None.")
    add("")

    add("## Distinct course codes")
    add("")
    add(", ".join(f"{code} ({n})" for code, n in sorted(codes.items())))
    add("")

    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        default=DEFAULT_SOURCE,
        help=f"source tree of PDFs/DOCX, relative to the repo root (default: {DEFAULT_SOURCE})",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=DEFAULT_OUT,
        help=f"markdown output tree, relative to the repo root (default: {DEFAULT_OUT})",
    )
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=REPO_ROOT,
        help="repo root that --source and --out are resolved against",
    )
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="OCR every page with Tesseract; only worth it if the report shows missing text layers",
    )
    parser.add_argument(
        "--report-only",
        action="store_true",
        help="rebuild report.md from an existing index.json without reconverting",
    )
    args = parser.parse_args()

    root = args.repo_root.resolve()
    source = args.source if args.source.is_absolute() else root / args.source
    out = args.out if args.out.is_absolute() else root / args.out

    if not source.is_dir():
        print(f"source not found: {args.source}", file=sys.stderr)
        return 1

    roots = Roots(
        source=source,
        out=out,
        rel_source=source.relative_to(root),
        rel_out=out.relative_to(root),
    )

    index_path = out / "index.json"
    totals = Totals()

    if args.report_only:
        if not index_path.is_file():
            print(f"no index to report on: {roots.rel_out}/index.json", file=sys.stderr)
            return 1
        stored = json.loads(index_path.read_text(encoding="utf-8"))
        totals.records = [Record.from_dict(d) for d in stored["records"]]
        used_ocr = stored.get("ocr", False)
    else:
        used_ocr = args.ocr
        files = source_files(source)
        if not files:
            print(f"no .pdf or .docx files under {roots.rel_source}", file=sys.stderr)
            return 1

        for i, path in enumerate(files, 1):
            print(f"[{i}/{len(files)}] {path.relative_to(source)}", file=sys.stderr)
            totals.records.append(convert_one(path, roots, use_ocr=args.ocr))

        out.mkdir(parents=True, exist_ok=True)
        offering_dicts = [r.offering for r in totals.of_kind("offering")]
        courses = build_courses(offering_dicts)
        index = {
            "source_root": roots.rel_source.as_posix(),
            "output_root": roots.rel_out.as_posix(),
            "ocr": args.ocr,
            "count": len(totals.records),
            "course_count": len(courses),
            # Two levels. A record is one file; a course groups the offering
            # records that share a `course_key`. Code alone is not identity.
            "courses": [c.to_dict() for c in courses],
            "title_review": [p.to_dict() for p in build_title_review(courses)],
            "records": [r.to_dict() for r in totals.records],
        }
        index_path.write_text(
            json.dumps(index, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )

    report = build_report(totals, roots.rel_source, roots.rel_out, use_ocr=used_ocr)
    (out / "report.md").write_text(report + "\n", encoding="utf-8")
    print(report)

    return 1 if any(not r.converted for r in totals.records) else 0


if __name__ == "__main__":
    raise SystemExit(main())
